from kivy.app import App
from kivy.uix.screenmanager import Screen
from kivy.metrics import dp
from kivy.properties import ObjectProperty, StringProperty
from kivymd.uix.datatables import MDDataTable

from kivymd.uix.boxlayout import MDBoxLayout

from kivymd.uix.dialog import MDDialog
from kivymd.uix.button import MDFlatButton

from modules.datamanager import Exam, Student, Grading, Subjects, SchoolCalendar
from modules.msgmanager import show_error, show_toast
from kivymd.uix.button import MDFloatingActionButtonSpeedDial
from kivymd.uix.menu import MDDropdownMenu

import docx, re, os
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT
        
from collections import Counter
from datetime import datetime, date

from operator import itemgetter
        
        
class ExamEditDialogContent(MDBoxLayout):
    pass
    

class ExamDataTable(MDDataTable):
    def __init__(self, **kwargs):
        super(ExamDataTable, self).__init__(**kwargs)
        
        
class ExamScreen(Screen):
    year = ObjectProperty()
    term = ObjectProperty()
    exam = ObjectProperty()
    form = ObjectProperty()
    
    dialog = None
    
    def __init__(self, **kwargs):
        super(ExamScreen, self).__init__(**kwargs)
        self.__adm_no = ""
       
    def on_pre_enter(self):
        self.update_exam_type()
       
    def on_pre_leave(self):
        self.ids.exam_container.clear_widgets()
           
    def open_print_menu(self):
        menu_items = [{
                       "viewclass": "OneLineListItem",
                       "text": "MARKS ENTRY LIST",
                       "height": dp(50),
                       "on_release": lambda: self.print_marks_list(),
                      }, 
                      {
                       "viewclass": "OneLineListItem",
                       "text": "EXAM ANALYSIS",
                       "height": dp(50),
                       "on_release": lambda: self.print_analysis_report(),
                      },
                      {
                      "viewclass": "OneLineListItem",
                      "text": "REPORT CARDS",
                      "height": dp(50),
                      "on_release": lambda: self.generate_and_print_reports(),
                      }
                      ]
        
        self.print_menu = MDDropdownMenu(
            caller=self.ids.print_btn,
            items=menu_items,
            position="bottom",
        )
        
        self.print_menu.open()
        
    def update_exam_type(self):
        self.ids.exam_container.clear_widgets()
        
        all_exams = []
        
        exm = Exam()
        exam_data = exm.get_exams(self.year.text, self.term.text, self.exam.text, self.form.text)
        
        if not exam_data:
            return
            
        # GRADING DATA 
     
        grading = Grading()
        
        grd_data = grading.get_grading_data()
        grading_data = [dict["GENERAL"] for dict in grd_data if "GENERAL" in dict][0]
        
        grd_data = grading.get_grading_data()
        science = [dict["SCIENCE"] for dict in grd_data if "SCIENCE" in dict][0]
        
        grd_data = grading.get_grading_data()
        language = [dict["LANGUAGES"] for dict in grd_data if "LANGUAGES" in dict][0]
        
        grd_data = grading.get_grading_data()
        humanity = [dict["HUMANITIES"] for dict in grd_data if "HUMANITIES" in dict][0]
       
        # ---------
            
        for exam in exam_data:
            adm_no = exam["adm"]
            
            marks = exam["marks"]
            
            std = Student()
            student_name = std.get_student_name(adm_no)
            
            eng = marks["ENG"]
            eng_g, eng_p = self.get_grade_and_points(language,self.wash_value(eng))
            
            kis = marks["KIS"]
            kis_g, kis_p = self.get_grade_and_points(language,self.wash_value(kis))
            
            mat = marks["MAT"]
            mat_g, mat_p = self.get_grade_and_points(science,self.wash_value(mat))
           
            bio = marks["BIO"]
            bio_g, bio_p = self.get_grade_and_points(science,self.wash_value(bio))
            
            phy = marks["PHY"]
            phy_g, phy_p = self.get_grade_and_points(science,self.wash_value(phy))
            
            che = marks["CHE"]
            che_g, che_p = self.get_grade_and_points(science,self.wash_value(che))
           
            his = marks["HIS"]
            his_g, his_p = self.get_grade_and_points(humanity,self.wash_value(his))
            
            geo = marks["GEO"]
            geo_g, geo_p = self.get_grade_and_points(humanity,self.wash_value(geo))
            
            cre = marks["CRE"]
            cre_g, cre_p = self.get_grade_and_points(humanity,self.wash_value(cre))
            
            agr = marks["AGR"]
            agr_g, agr_p = self.get_grade_and_points(humanity,self.wash_value(agr))
            
            com = marks["COM"]
            com_g, com_p = self.get_grade_and_points(humanity,self.wash_value(com))
            
            bus = marks["BUS"]
            bus_g, bus_p = self.get_grade_and_points(humanity,self.wash_value(bus))
          
            all_exams.append((f"[size=24]{adm_no}[/size]", 
                              f"[size=24]{student_name}[/size]", 
                              f"[size=24]{eng}{eng_g}[/size]", 
                              f"[size=24]{kis}{kis_g}[/size]",
                              f"[size=24]{mat}{mat_g}[/size]", 
                              f"[size=24]{bio}{bio_g}[/size]", 
                              f"[size=24]{phy}{phy_g}[/size]", 
                              f"[size=24]{che}{che_g}[/size]",
                              f"[size=24]{his}{his_g}[/size]", 
                              f"[size=24]{geo}{geo_g}[/size]", 
                              f"[size=24]{agr}{agr_g}[/size]", 
                              f"[size=24]{com}{com_g}[/size]", 
                              f"[size=24]{bus}{bus_g}[/size]",
                              f"[size=24]{cre}{cre_g}[/size]", 
                              ))
           
        data_tables = ExamDataTable(
            column_data=[
                ("[size=24]ADM[/size]", dp(15)),
                ("[size=24]STUDENT NAME[/size]", dp(30)),
                ("[size=24]ENG[/size]", dp(15)),
                ("[size=24]KIS[/size]", dp(15)),
                ("[size=24]MAT[/size]", dp(15)),
                ("[size=24]BIO[/size]", dp(15)),
                ("[size=24]PHY[/size]", dp(15)),
                ("[size=24]CHE[/size]", dp(15)),
                ("[size=24]HIS[/size]", dp(15)),
                ("[size=24]GEO[/size]", dp(15)),
                ("[size=24]AGR[/size]", dp(15)),
                ("[size=24]COM[/size]", dp(15)),
                ("[size=24]BUS[/size]", dp(15)),
                ("[size=24]CRE[/size]", dp(15)),
                ],
            
            row_data = [row for row in all_exams],
        )
            
        data_tables.bind(on_row_press=self.view_exam)
        self.ids.exam_container.add_widget(data_tables)
            
   
    def view_exam(self, table_view, cell_row):
        start_index, end_index = cell_row.table.recycle_data[cell_row.index]["range"]
        cell_data = cell_row.table.recycle_data[start_index]["text"]
        
        self.__adm_no = re.findall("[\d]+", cell_data)[1]
        
        std = Student()
        sname = std.get_student_name(self.__adm_no)
        
        exm = Exam()
        exam = exm.get_student_marks(self.__adm_no, self.year.text, self.term.text, self.exam.text, self.form.text)
        
        marks = exam["marks"]
        
        eng = marks["ENG"]
        kis = marks["KIS"]
        mat = marks["MAT"]
            
        bio = marks["BIO"]
        phy = marks["PHY"]
        che = marks["CHE"]
            
        his = marks["HIS"]
        geo = marks["GEO"]
        cre = marks["CRE"]
           
        agr = marks["AGR"]
        com = marks["COM"]
        bus = marks["BUS"]
        
        if not self.dialog:
            self.dialog = MDDialog(
                title = f"{sname} | {self.__adm_no}",
                type = "custom",
                content_cls = ExamEditDialogContent(),
                auto_dismiss = False,
                buttons = [
                    MDFlatButton(
                        text = "CANCEL",
                        theme_text_color ="Custom",
                        on_release = lambda _: self.close_dialog(),
                        ),
                        
                    MDFlatButton(
                        text = "DELETE",
                        theme_text_color ="Custom",
                        on_release = lambda _: self.delete_student_exam(),
                        ),
               
                    MDFlatButton(
                        text = "SAVE",
                        theme_text_color = "Custom",
                        on_release = lambda _: self.update_exam(),
                        ),
                    ], 
                )
                
            self.dialog.content_cls.ids.eng.text = eng
            self.dialog.content_cls.ids.kis.text = kis
            self.dialog.content_cls.ids.mat.text = mat
            
            self.dialog.content_cls.ids.bio.text = bio
            self.dialog.content_cls.ids.phy.text = phy
            self.dialog.content_cls.ids.che.text = che
            
            self.dialog.content_cls.ids.his.text = his
            self.dialog.content_cls.ids.geo.text = geo
            self.dialog.content_cls.ids.cre.text = cre
                
            self.dialog.content_cls.ids.agr.text = agr
            self.dialog.content_cls.ids.com.text = com
            self.dialog.content_cls.ids.bus.text = bus
            
            
        self.dialog.open()
        
        return
        
    def delete_student_exam(self):
        self.dialog.dismiss()
        
        self.dialog_2 = MDDialog(
                title = "Delete Marks", 
                text = f"Delete '{self.term.text}' marks for ADM No. {self.__adm_no}?",
                type = "custom",
                auto_dismiss = False,
                        
                buttons = [
                    MDFlatButton(
                        text = "CANCEL",
                        theme_text_color = "Custom",
                        text_color = "#008080",
                        pos_hint = {'center_x': .5, 'center_y': .5},
                        on_release = lambda _: self.dialog_2.dismiss(),
                        ),
                        
                    MDFlatButton(
                        text = "DELETE",
                        theme_text_color = "Custom",
                        text_color = "#FF0000",
                        pos_hint = {'center_x': .5, 'center_y': .5},
                        on_release = lambda _: self.delete_exam(),
                        ),
                ],
            )
            
        self.dialog_2.open()
       
        return
        
    def delete_exam(self):
        self.dialog_2.dismiss()
            
        exm = Exam()
        stdexam = exm.get_student_exam(self.__adm_no, self.year.text, self.term.text, self.exam.text, self.form.text)
        std_id = stdexam["_id"]
        
        exm.delete_exam(std_id)
        
        show_toast(f"Deleted '{self.term.text}' marks for ADM No. {self.__adm_no}")
              
        return
        
    def close_dialog(self):
        self.dialog.dismiss()
        self.dialog = None
        
        return
        
    def update_exam(self):
        eng = self.dialog.content_cls.ids.eng.text
        kis = self.dialog.content_cls.ids.kis.text 
        mat = self.dialog.content_cls.ids.mat.text 
            
        bio = self.dialog.content_cls.ids.bio.text 
        phy = self.dialog.content_cls.ids.phy.text 
        che = self.dialog.content_cls.ids.che.text 
            
        his = self.dialog.content_cls.ids.his.text 
        geo = self.dialog.content_cls.ids.geo.text 
        cre = self.dialog.content_cls.ids.cre.text 
                
        agr = self.dialog.content_cls.ids.agr.text
        com = self.dialog.content_cls.ids.com.text 
        bus = self.dialog.content_cls.ids.bus.text
        
        self.close_dialog()
        
        grading = Grading()
        
        grd_data = grading.get_grading_data()
        sci = [dict["SCIENCE"] for dict in grd_data if "SCIENCE" in dict][0]
        
        grd_data = grading.get_grading_data()
        lang = [dict["LANGUAGES"] for dict in grd_data if "LANGUAGES" in dict][0]
        
        grd_data = grading.get_grading_data()
        hum = [dict["HUMANITIES"] for dict in grd_data if "HUMANITIES" in dict][0]
        
        # EXAM OBJ
        
        exm = Exam()
        stdexam = exm.get_student_exam(self.__adm_no, self.year.text, self.term.text, self.exam.text, self.form.text)
        std_id = stdexam["_id"]
        
        edited_marks = {"ENG": eng, 
                      "KIS": kis, 
                      "MAT": mat,
                      "BIO": bio,
                      "PHY": phy,
                      "CHE": che,
                      "HIS": his,
                      "GEO": geo,
                      "CRE": cre,
                      "AGR": agr,
                      "COM": com,
                      "BUS": bus
            }
            
        #exm.update_marks(self.__adm_no, self.year.text, self.term.text, self.exam.text, self.form.text, edited_marks)
        exm.update_marks(std_id, edited_marks)
        
        show_toast(f"Updated marks for ADM No. {self.__adm_no}")
                
        return
        
    def process_exam_data(self, exam_type):
        student_exams_data = []
        
        exm = Exam()
        exam_data = exm.get_exams(self.year.text, self.term.text, exam_type, self.form.text)
        
        # GRADING DATA 
        
        grd_data = []
        
        grading = Grading()
        
        for grd in grading.get_grading_data():
            grd_data.append(grd)
        
        grading_data = [dict["GENERAL"] for dict in grd_data if "GENERAL" in dict][0]
        
        science = [dict["SCIENCE"] for dict in grd_data if "SCIENCE" in dict][0]
        language = [dict["LANGUAGES"] for dict in grd_data if "LANGUAGES" in dict][0]
        humanity = [dict["HUMANITIES"] for dict in grd_data if "HUMANITIES" in dict][0]
       
        # ---------
        
        for exam in exam_data:
            adm_no = exam["adm"]
            
            std = Student()
            student_name = std.get_student_name(adm_no)
            
            marks = exam["marks"]
            
            eng = marks["ENG"]
            eng_g, eng_p = self.get_grade_and_points(language,self.wash_value(eng))
            
            kis = marks["KIS"]
            kis_g, kis_p = self.get_grade_and_points(language,self.wash_value(kis))
            
            mat = marks["MAT"]
            mat_g, mat_p = self.get_grade_and_points(science,self.wash_value(mat))
           
            bio = marks["BIO"]
            bio_g, bio_p = self.get_grade_and_points(science,self.wash_value(bio))
            
            phy = marks["PHY"]
            phy_g, phy_p = self.get_grade_and_points(science,self.wash_value(phy))
            
            che = marks["CHE"]
            che_g, che_p = self.get_grade_and_points(science,self.wash_value(che))
           
            his = marks["HIS"]
            his_g, his_p = self.get_grade_and_points(humanity,self.wash_value(his))
            
            geo = marks["GEO"]
            geo_g, geo_p = self.get_grade_and_points(humanity,self.wash_value(geo))
            
            cre = marks["CRE"]
            cre_g, cre_p = self.get_grade_and_points(humanity,self.wash_value(cre))
            
            agr = marks["AGR"]
            agr_g, agr_p = self.get_grade_and_points(humanity,self.wash_value(agr))
            
            com = marks["COM"]
            com_g, com_p = self.get_grade_and_points(humanity,self.wash_value(com))
            
            bus = marks["BUS"]
            bus_g, bus_p = self.get_grade_and_points(humanity,self.wash_value(bus))
            
            student_exam_data = [[adm_no, student_name],
                              [(self.wash_value(eng), eng_g, eng_p),
                               (self.wash_value(kis), kis_g, kis_p),
                               (self.wash_value(mat), mat_g, mat_p),
                               (self.wash_value(bio), bio_g, bio_p), 
                               (self.wash_value(phy), phy_g, phy_p),
                               (self.wash_value(che), che_g, che_p),
                               (self.wash_value(his), his_g, his_p),
                               (self.wash_value(geo), geo_g, geo_p),
                               (self.wash_value(cre), cre_g, cre_p),
                               (self.wash_value(agr), agr_g, agr_p),
                               (self.wash_value(com), com_g, com_p),
                               (self.wash_value(bus), bus_g, bus_p),
                               ]
                              ]    
                                             
            if self.form.text == "FORM 1" or self.form.text == "FORM 2":
                #---
                exams_did = [marks[0] for marks in student_exam_data[1] if marks[0] != -1]
                total_exams_did = len(exams_did)  
                total_marks = sum(exams_did)
                total_marks_out_of = total_exams_did * 100
                avg_marks = int(total_marks / total_exams_did)
                    
                mean_grade = self.get_f1f2_mean_grade(avg_marks)
                
                if total_exams_did < 8:
                    mean_grade = "X"
                        
                student_exam_data.append([total_marks, mean_grade])
                
            if self.form.text == "FORM 3" or self.form.text == "FORM 4": 
                #---
                exams_did = [marks[0] for marks in student_exam_data[1] if marks[0] != -1]
                total_exams_did = len(exams_did) 
                total_marks = sum(exams_did)
                total_points = sum([points[2] for points in student_exam_data[1]])
                total_points_out_of = 84
            
                if total_exams_did < 7:
                    mean_grade = "X"
                        
                elif total_exams_did == 7:
                    mean_grade = self.get_f3f4_mean_grade(total_points)
                    
                elif total_exams_did >= 8:
                    best_points = []
                       
                    math = [points[2] for points in student_exam_data[1]][2]
                    eng, kis = [points[2] for points in student_exam_data[1]][0:2]
                    
                    best_lang = self.get_one_best_point_among_two(eng, kis)
                    best5 = self.get_five_best_points_among_many([points[2] for points in student_exam_data[1]][3:])
                    
                    best_points.append(math)
                    best_points.append(best_lang)
                    best_points.extend(best5)
                    
                    total_points = sum(best_points)
                    mean_grade = self.get_f3f4_mean_grade(total_points)
                        
                else:
                    continue
                
                student_exam_data.append([total_points, mean_grade])
            
            student_exams_data.append(student_exam_data)
            
        return student_exams_data
        
    def generate_and_print_reports(self):
        exams = ["OPENER","MID TERM","END TERM"]
        
        term_exams = {}
        
        if self.exam.text == exams[0]:
            opener_data = self.process_exam_data(exams[0])
            term_exams[exams[0]] = opener_data
            
        elif self.exam.text == exams[1]:
            opener_data = self.process_exam_data(exams[0])
            term_exams[exams[0]] = opener_data
            
            midterm_data = self.process_exam_data(exams[1])
            term_exams[exams[1]] = midterm_data
            
        elif self.exam.text == exams[2]:
            opener_data = self.process_exam_data(exams[0])
            term_exams[exams[0]] = opener_data
            
            midterm_data = self.process_exam_data(exams[1])
            term_exams[exams[1]] = midterm_data
            
            endterm_data = self.process_exam_data(exams[2])
            term_exams[exams[2]] = endterm_data
            
        else:
            pass
            
        with open(f"exam_report_cards_test_{datetime.now().hour}_{datetime.now().minute}_{datetime.now().second}.txt", "w") as examout:
            for k,v in term_exams.items():
                examout.write(f"{str(k)}\n\n")
                examout.write("=="*10 + "\n")
                examout.write(f"{str(v)}\n\n")
                examout.write("×××××\n\n")
                
        f = open(f"exam_dump_test_{datetime.now().hour}_{datetime.now().minute}_{datetime.now().second}.txt", "w")
        f.write(str(term_exams))
        f.close()
        
        show_toast(f"Successfully written data to file.")
   
    def print_report_cards(self, data_items, marks_list):
        from modules.printermanager import SchoolDocument
        
        app_path = App.get_running_app().download_path
        filename = f"report_cards_{re.sub(' ', '_', self.form.text.lower())}_{re.sub(' ', '_', self.term.text.lower())}_{re.sub(' ', '_', self.exam.text.lower())}_{self.year.text}.docx"
        
        ### CALENDAR DATES ###
        
        sch_calendas = SchoolCalendar()
        calendas = [cal for cal in sch_calendas.get_schcalendars()]
        
        try:
            if self.exam.text == "MID TERM":
                sch_closing_date = [data[self.year.text][self.term.text]["half_term"]["start"] for data in calendas if self.year.text in data][0]
                sch_opening_date = [data[self.year.text][self.term.text]["half_term"]["end"] for data in calendas if self.year.text in data][0]
            
            elif self.exam.text == "END TERM":
                sch_closing_date = [data[self.year.text][self.term.text]["holiday"]["start"] for data in calendas if self.year.text in data][0]
                sch_opening_date = [data[self.year.text][self.term.text]["holiday"]["end"] for data in calendas if self.year.text in data][0]
            
            else:
                show_error("Print Exam", "Sorry! No transcripts for opener exams!")
                
                return
                
            sch_closing_date = datetime.strptime(sch_closing_date, "%Y-%m-%d").strftime("%A, %d %B %Y")
            sch_opening_date = datetime.strptime(sch_opening_date, "%Y-%m-%d").strftime("%A, %d %B %Y")
            
        except:
            sch_closing_date = "00-00-0000"
            sch_opening_date = "00-00-0000"
            
        marks_list = sorted(marks_list, reverse=True, key=itemgetter(0))
        
        out_of = len(marks_list)
        
        ### DOCUMENT WRITER ###
        
        docwriter = SchoolDocument("", 11)
        docwriter.write_headers_and_footers()
            
        # ---------
        
        for exam_data, marks_data, totals_data, stats in data_items:
            
            # Header Data
            
            adm_no = exam_data[0]
            student_name = exam_data[1]
            year = exam_data[2]
            term = exam_data[3]
            exam = exam_data[4]
            form = exam_data[5]
            
            grade = stats[0]
            marks = stats[1]
            
            hdtrs_comment, clstrs_comment = self.get_teachers_comments(grade)
            
            # STUDENT CLASS POSITION
            
            position = marks_list.index((marks, str(adm_no))) + 1
           
            
            ## WRITING TO DOCWRITER ##
        
            docwriter.write_headings()
            docwriter.write_title(f"STUDENT REPORT CARD")
        
            doc = docwriter.write_next()
            
            ## ----------- ##
            
            name_para = doc.add_paragraph()
            paragraph_format = name_para.paragraph_format
            paragraph_format.line_spacing = Pt(15)

            name_para.add_run("NAME ")

            std_name = name_para.add_run(f"     {student_name.upper()}      ")
            std_name.underline = True
            std_name.underline = WD_UNDERLINE.DOTTED
            std_name.bold = True
            std_name.font.name = "Lucida Calligraphy"

            name_para.add_run("ADM NO. ")

            adm_no = name_para.add_run(f"      {adm_no}      ")
            adm_no.underline = True
            adm_no.underline = WD_UNDERLINE.DOTTED
            adm_no.bold = True
            adm_no.font.name = "Lucida Calligraphy"

            name_para.add_run("FORM ")

            form = name_para.add_run(f"      {form}     ")
            form.underline = True
            form.underline = WD_UNDERLINE.DOTTED
            form.bold = True

            name_para.add_run("TERM ")

            term = name_para.add_run(f"     {term}     ")
            term.underline = True
            term.underline = WD_UNDERLINE.DOTTED
            term.bold = True
            term.font.name = "Lucida Calligraphy"
    
            name_para.add_run("EXAM")

            exam = name_para.add_run(f"     {exam}     ")
            exam.underline = True
            exam.underline = WD_UNDERLINE.DOTTED
            exam.bold = True
            exam.font.name = "Lucida Calligraphy"

            # YEAR
            
            name_para.add_run("YEAR ")

            year = name_para.add_run(f"     {year}     _")
            year.underline = True
            year.underline = WD_UNDERLINE.DOTTED
            year.bold = True
            year.font.name = "Lucida Calligraphy"
           
            position_para = doc.add_paragraph()
            paragraph_format = position_para.paragraph_format
            paragraph_format.line_spacing = Pt(18)
            
            # POSITION
            
            position_para.add_run("MEAN GRADE ")

            grade = position_para.add_run(f"         {grade}        ")
            grade.underline = True
            grade.underline = WD_UNDERLINE.DOTTED
            grade.bold = True
            grade.font.name = "Lucida Calligraphy"
            
            position_para.add_run("POSITION")
            
            pstn = position_para.add_run(f"          {position}          ")
            pstn.underline = True
            pstn.underline = WD_UNDERLINE.DOTTED
            pstn.bold = True
            pstn.font.name = "Lucida Calligraphy"
            
            # OUT OF
            
            position_para.add_run("OUT OF")
            
            outof = position_para.add_run(f"          {out_of}         _")
            outof.underline = True
            outof.underline = WD_UNDERLINE.DOTTED
            outof.bold = True
            outof.font.name = "Lucida Calligraphy"
    
            # ADD MARKS 
    
            table1 = doc.add_table(rows = 1, cols = 7)
            table1.style = doc.styles["Table Grid"]

            table1.columns[0].width = Cm(1.5)
            table1.columns[1].width = Cm(6.5)
            table1.columns[2].width = Cm(3.5)
            table1.columns[3].width = Cm(3.5)
            table1.columns[4].width = Cm(3.5)
            table1.columns[5].width = Cm(8.0)
            table1.columns[6].width = Cm(4.5)
    
            cells = table1.rows[0].cells
    
            cells[0].text = "CODE"
            cells[0].width = Cm(1.5)
    
            cells[1].text = "SUBJECTS"
            cells[1].width = Cm(6.5)
    
            cells[2].text = "EXAM(%)"
            cells[2].width = Cm(3.5)
    
            cells[3].text = "GRADE"
            cells[3].width = Cm(3.5)
    
            cells[4].text = "POINTS"
            cells[4].width = Cm(3.5)
    
            cells[5].text = "REMARKS"
            cells[5].width = Cm(8.0)
    
            cells[6].text = "TEACHER"
            cells[6].width = Cm(4.5)
          
            for row in marks_data:
                code, subject, marks, grd, points, remark, sign = row
        
                row_cells = table1.add_row().cells
        
                row_cells[0].text = str(code)
                row_cells[1].text = subject
                row_cells[2].text = str(marks) if marks != -1 else "" 
                row_cells[3].text = grd
                row_cells[4].text = str(points) if points != 0 else ""
                row_cells[5].text = remark
                row_cells[6].text = sign
        
            for row in totals_data:
                rname, _, total_marks, _,total_points, _, _ = row
        
                row_cells = table1.add_row().cells
        
                row_cells[0].merge(row_cells[1])
        
                row_cells[0].text = rname
                row_cells[1].text = ""
                row_cells[2].text = str(total_marks)
                row_cells[3].text = ""
                row_cells[4].text = str(total_points)
                row_cells[5].text = ""
                row_cells[6].text = ""
        
            remarks_heading = doc.add_heading()
            remarks_ = remarks_heading.add_run("REMARKS", 0)
            remarks_.font.name = "Times New Roman"
            remarks_.font.color.rgb = RGBColor.from_string("0047AB")
          
            ct_trs_remark = doc.add_paragraph("CLASS TEACHER'S COMMENTS:")
            ct_comment = ct_trs_remark.add_run(f"     {clstrs_comment}   _\n")
            ct_comment.underline = True
            ct_comment.underline = WD_UNDERLINE.DOTTED
            ct_comment.bold = True
            ct_comment.font.name = "Lucida Calligraphy"
            
            ct_trs_remark.add_run(f"Date: ")
            ct_trs_remark.add_run("."*50)
            ct_trs_remark.add_run(" "*10)
            ct_trs_remark.add_run("Signature: ")
            ct_trs_remark.add_run("."*50)
    
            prp_trs_remark = doc.add_paragraph("PRINCIPAL'S COMMENTS:")
            prp_comment = prp_trs_remark.add_run(f"     {hdtrs_comment}    _\n")
            prp_comment.underline = True
            prp_comment.underline = WD_UNDERLINE.DOTTED
            prp_comment.bold = True
            prp_comment.font.name = "Lucida Calligraphy"
            
            prp_trs_remark.add_run("Date: ")
            prp_trs_remark.add_run("."*50)
            prp_trs_remark.add_run(" "*10)
            prp_trs_remark.add_run("Signature: ")
            prp_trs_remark.add_run("."*50)
       
            parent_seen = doc.add_paragraph("Parent / Guardian's Signature:  ")
            parent_seen.add_run("."*80)
            
            next_term = doc.add_paragraph("School has closed today on ")
            term_date_closing = next_term.add_run(f"       {sch_closing_date}         _")
            term_date_closing.underline = True
            term_date_closing.underline = WD_UNDERLINE.DOTTED
            term_date_closing.bold = True
            term_date_closing.font.name = "Lucida Calligraphy"
            
            next_term.add_run(" and reopens next time on ") 
            term_date_opening = next_term.add_run(f"       {sch_opening_date}         _")
            term_date_opening.underline = True
            term_date_opening.underline = WD_UNDERLINE.DOTTED
            term_date_opening.bold = True
            term_date_opening.font.name = "Lucida Calligraphy"
            
            doc.add_page_break()
            
        docwriter.save_document(os.path.join(app_path, filename))
        
        show_toast("Printed transcripts to 'Download' folder.")
        
        return
    
    def print_marks_list(self):
        self.print_menu.dismiss()
        
        from modules.printermanager import SchoolDocument
        
        app_path = App.get_running_app().download_path
        filename = f"marks_entry_list_{re.sub(' ', '_', self.form.text.lower())}_{re.sub(' ', '_', self.term.text.lower())}_{re.sub(' ', '_', self.exam.text.lower())}_{self.year.text}.docx"
        
        student_list = []
        
        std_obj = Student()
        students_data = std_obj.get_students(self.form.text)
        
        for student in students_data:
            name = student["full_name"]
            adm = student["_id"]
                
            student_list.append((adm, name))
            
        ### DOCUMENT WRITER ###
        
        docwriter = SchoolDocument("", 11)
        docwriter.set_landscape_mode()
        docwriter.write_headers_and_footers()
        docwriter.write_headings()
        docwriter.write_title(f"Marks Entry List, {self.form.text}, {self.exam.text} exam, {self.term.text}, {self.year.text}")
    
        doc = docwriter.write_next()
        
        table1 = doc.add_table(rows = 1, cols = 17)
        table1.style = doc.styles["Light Grid Accent 1"]

        table1.columns[0].width = Cm(2.5)
        table1.columns[1].width = Cm(12.5)
        table1.columns[2].width = Cm(2.5)
        table1.columns[3].width = Cm(2.5)
        table1.columns[4].width = Cm(2.5)
        table1.columns[5].width = Cm(2.5)
        table1.columns[6].width = Cm(2.5)
        table1.columns[7].width = Cm(2.5)
        table1.columns[8].width = Cm(2.5)
        table1.columns[9].width = Cm(2.5)
        table1.columns[10].width = Cm(2.5)
        table1.columns[11].width = Cm(2.5)
        table1.columns[12].width = Cm(2.5)
        table1.columns[13].width = Cm(2.5)
        table1.columns[14].width = Cm(2.5)
        table1.columns[15].width = Cm(2.5)
        table1.columns[16].width = Cm(2.5)
        
        cells = table1.rows[0].cells
        
        cells0100 = cells[0].paragraphs[0]
        cells0100text = cells0100.add_run("ADM")
        cells0100text.bold = True
        cells[0].width = Cm(2.5)
        
        cells0101 = cells[1].paragraphs[0]
        cells0101text = cells0101.add_run("NAME")
        cells0101text.bold = True
        cells[1].width = Cm(12.5)

        cells0102 = cells[2].paragraphs[0]
        cells0102text = cells0102.add_run("ENG")
        cells0102text.bold = True
        cells[2].width = Cm(2.5)

        cells0103 = cells[3].paragraphs[0]
        cells0103text = cells0103.add_run("KIS")
        cells0103text.bold = True
        cells[3].width = Cm(2.5)

        cells0104 = cells[4].paragraphs[0]
        cells0104text = cells0104.add_run("MAT")
        cells0104text.bold = True
        cells[4].width = Cm(2.5)

        cells0105 = cells[5].paragraphs[0]
        cells0105text = cells0105.add_run("BIO")
        cells0105text.bold = True
        cells[5].width = Cm(2.5)

        cells0106 = cells[6].paragraphs[0]
        cells0106text = cells0106.add_run("PHY")
        cells0106text.bold = True
        cells[6].width = Cm(2.5)

        cells0107 = cells[7].paragraphs[0]
        cells0107text = cells0107.add_run("CHE")
        cells0107text.bold = True
        cells[7].width = Cm(2.5)

        cells0108 = cells[8].paragraphs[0]
        cells0108text = cells0108.add_run("HIS")
        cells0108text.bold = True
        cells[8].width = Cm(2.5)

        cells0109 = cells[9].paragraphs[0]
        cells0109text = cells0109.add_run("GEO")
        cells0109text.bold = True
        cells[9].width = Cm(2.5)

        cells0110 = cells[10].paragraphs[0]
        cells0110text = cells0110.add_run("CRE")
        cells0110text.bold = True
        cells[10].width = Cm(2.5)
        
        cells0111 = cells[11].paragraphs[0]
        cells0111text = cells0111.add_run("AGR")
        cells0111text.bold = True
        cells[11].width = Cm(2.5)

        cells0112 = cells[12].paragraphs[0]
        cells0112text = cells0112.add_run("COM")
        cells0112text.bold = True
        cells[12].width = Cm(2.5)
        
        cells0113 = cells[13].paragraphs[0]
        cells0113text = cells0113.add_run("BUS")
        cells0113text.bold = True
        cells[13].width = Cm(2.5)
        
        cells0114 = cells[14].paragraphs[0]
        cells0114text = cells0114.add_run("TOT")
        cells0114text.bold = True
        cells[14].width = Cm(2.5)
        
        cells0115 = cells[15].paragraphs[0]
        cells0115text = cells0115.add_run("GRD")
        cells0115text.bold = True
        cells[15].width = Cm(2.5)
        
        cells0116 = cells[16].paragraphs[0]
        cells0116text = cells0116.add_run("PST")
        cells0116text.bold = True
        cells[16].width = Cm(2.5)
        
        student_list = sorted(student_list, reverse=False, key=itemgetter(0))
        
        for adm, student in student_list:
            row_cells = table1.add_row().cells
    
            row_cells[0].text = str(adm)
            row_cells[1].text = str(student)
            row_cells[2].text = ""
            row_cells[3].text = ""
            row_cells[4].text = ""
            row_cells[5].text = ""
            row_cells[6].text = ""
            row_cells[7].text = ""
            row_cells[8].text = ""
            row_cells[9].text = ""
            row_cells[10].text = ""
            row_cells[11].text = ""
            row_cells[12].text = ""
            row_cells[13].text = ""
            row_cells[14].text = ""
            row_cells[15].text = ""
            row_cells[16].text = ""
      
        docwriter.save_document(os.path.join(app_path, filename))
        
        show_toast("Marks list saved to 'Download' folder.")
        
    def print_analysis_report(self):
        self.print_menu.dismiss()
        
        from modules.printermanager import SchoolDocument
        
        app_path = App.get_running_app().download_path
        filename = f"exam_analysis_{re.sub(' ', '_', self.form.text.lower())}_{re.sub(' ', '_', self.term.text.lower())}_{re.sub(' ', '_', self.exam.text.lower())}_{self.year.text}.docx"
        
        exm = Exam()
        exam_data = exm.get_exams(self.year.text, self.term.text, self.exam.text, self.form.text)
        
        if not exam_data:
            show_error("Print Report", "There are no exams to print.")
            return
       
        # GRADING DATA 
     
        grading = Grading()
       
        grd_data = grading.get_grading_data()
        grading_data = [dict["GENERAL"] for dict in grd_data if "GENERAL" in dict][0]
        
        grd_data = grading.get_grading_data()
        science = [dict["SCIENCE"] for dict in grd_data if "SCIENCE" in dict][0]
        
        grd_data = grading.get_grading_data()
        language = [dict["LANGUAGES"] for dict in grd_data if "LANGUAGES" in dict][0]
        
        grd_data = grading.get_grading_data()
        humanity = [dict["HUMANITIES"] for dict in grd_data if "HUMANITIES" in dict][0]
       
        # ---------
        
        student_marks = []
         
        for exam in exam_data:
            adm_no = exam["adm"]
            
            marks = exam["marks"]
            
            std = Student()
            student_name = std.get_student_name(adm_no)
            
            eng = marks["ENG"]
            eng_g, eng_p = self.get_grade_and_points(language,self.wash_value(eng))
            
            kis = marks["KIS"]
            kis_g, kis_p = self.get_grade_and_points(language,self.wash_value(kis))
            
            mat = marks["MAT"]
            mat_g, mat_p = self.get_grade_and_points(science,self.wash_value(mat))
           
            bio = marks["BIO"]
            bio_g, bio_p = self.get_grade_and_points(science,self.wash_value(bio))
            
            phy = marks["PHY"]
            phy_g, phy_p = self.get_grade_and_points(science,self.wash_value(phy))
            
            che = marks["CHE"]
            che_g, che_p = self.get_grade_and_points(science,self.wash_value(che))
           
            his = marks["HIS"]
            his_g, his_p = self.get_grade_and_points(humanity,self.wash_value(his))
            
            geo = marks["GEO"]
            geo_g, geo_p = self.get_grade_and_points(humanity,self.wash_value(geo))
            
            cre = marks["CRE"]
            cre_g, cre_p = self.get_grade_and_points(humanity,self.wash_value(cre))
            
            agr = marks["AGR"]
            agr_g, agr_p = self.get_grade_and_points(humanity,self.wash_value(agr))
            
            com = marks["COM"]
            com_g, com_p = self.get_grade_and_points(humanity,self.wash_value(com))
            
            bus = marks["BUS"]
            bus_g, bus_p = self.get_grade_and_points(humanity,self.wash_value(bus))
         
            student_marks.append([adm_no, 
                                  student_name, 
                                  [self.wash_value(eng), eng_g, eng_p], 
                                  [self.wash_value(kis), kis_g, kis_p],
                                  [self.wash_value(mat), mat_g, mat_p],
                                  [self.wash_value(bio), bio_g, bio_p], 
                                  [self.wash_value(phy), phy_g, phy_p],
                                  [self.wash_value(che), che_g, che_p],
                                  [self.wash_value(his), his_g, his_p],
                                  [self.wash_value(geo), geo_g, geo_p], 
                                  [self.wash_value(cre), cre_g, cre_p],
                                  [self.wash_value(agr), agr_g, agr_p], 
                                  [self.wash_value(com), com_g, com_p],
                                  [self.wash_value(bus), bus_g, bus_p]
                                  ]
                                 ) # Removed IRE and HRE from list. Its not available in this school.
                                 
            if self.form.text == "FORM 1" or self.form.text == "FORM 2": 
                for row in student_marks:
                    exams_did = [value[0] for value in row[2:] if isinstance(value, list) and isinstance(value[0], int) and value[0] != -1]
                    total_exams_did = len(exams_did)
                    
                    total_marks = sum(exams_did)
                    avg_marks = int(total_marks / total_exams_did)
                    
                    mean_grade = self.get_f1f2_mean_grade(avg_marks)
                
                    if total_exams_did < 8:
                        mean_grade = "X"
                        
                    row.append(total_marks)
                    row.append(mean_grade)
              
            if self.form.text == "FORM 3" or self.form.text == "FORM 4": 
                for row in student_marks:
                    exams_did = [value[0] for value in row[2:] if isinstance(value, list) and isinstance(value[0], int) and value[0] != -1]
                    total_exams_did = len(exams_did)
                    
                    total_points = sum([value[2] for value in row[2:] if isinstance(value, list) and isinstance(value[2], int)])
                    
                    if total_exams_did < 7:
                        mean_grade = "X"
                        
                    elif total_exams_did == 7:
                        mean_grade = self.get_f3f4_mean_grade(total_points)
                    
                    elif total_exams_did >= 8:
                        best_points = []
                       
                        math = row[4][2]
                        eng, kis = [row[2][2], row[3][2]]
                    
                        best_lang = self.get_one_best_point_among_two(eng, kis)
                        best5 = self.get_five_best_points_among_many([value[2] for value in row[5:] if isinstance(value, list) and isinstance(value[2], int)])
                    
                        best_points.append(math)
                        best_points.append(best_lang)
                        best_points.extend(best5)
                    
                        total_points = sum(best_points)
                        mean_grade = self.get_f3f4_mean_grade(total_points)
                        
                    else:
                        continue
                    
                    row.append(total_points)
                    row.append(mean_grade)
                     
            student_marks = sorted(student_marks, reverse=True, key=itemgetter(14))
       

            ### DOCUMENT WRITER ###
        
            docwriter = SchoolDocument("", 11)
            docwriter.set_landscape_mode()
            docwriter.write_headers_and_footers()
            docwriter.write_headings()
            docwriter.write_title(f"Exam Analysis Report, {self.form.text}, {self.exam.text} exam, {self.term.text}, {self.year.text}")
        
            doc = docwriter.write_next()
            
            gen_para = doc.add_paragraph()
            gen_para.paragraph_format.space_before = Inches(0)
            
            gen = gen_para.add_run("GENERAL ANALYSIS")
            gen.underline = True
            gen.font.color.rgb = RGBColor.from_string("FF0000")
           
            table1 = doc.add_table(rows = 1, cols = 17)
            table1.style = doc.styles["Light Grid Accent 1"]

            table1.columns[0].width = Cm(2.5)
            table1.columns[1].width = Cm(12.5)
            table1.columns[2].width = Cm(2.5)
            table1.columns[3].width = Cm(2.5)
            table1.columns[4].width = Cm(2.5)
            table1.columns[5].width = Cm(2.5)
            table1.columns[6].width = Cm(2.5)
            table1.columns[7].width = Cm(2.5)
            table1.columns[8].width = Cm(2.5)
            table1.columns[9].width = Cm(2.5)
            table1.columns[10].width = Cm(2.5)
            table1.columns[11].width = Cm(2.5)
            table1.columns[12].width = Cm(2.5)
            table1.columns[13].width = Cm(2.5)
            table1.columns[14].width = Cm(2.5)
            table1.columns[15].width = Cm(2.5)
            table1.columns[16].width = Cm(2.5)
            
            cells = table1.rows[0].cells
            
            cells0100 = cells[0].paragraphs[0]
            cells0100text = cells0100.add_run("ADM")
            cells0100text.bold = True
            cells[0].width = Cm(2.5)
            
            cells0101 = cells[1].paragraphs[0]
            cells0101text = cells0101.add_run("NAME")
            cells0101text.bold = True
            cells[1].width = Cm(12.5)
    
            cells0102 = cells[2].paragraphs[0]
            cells0102text = cells0102.add_run("ENG")
            cells0102text.bold = True
            cells[2].width = Cm(2.5)
    
            cells0103 = cells[3].paragraphs[0]
            cells0103text = cells0103.add_run("KIS")
            cells0103text.bold = True
            cells[3].width = Cm(2.5)
    
            cells0104 = cells[4].paragraphs[0]
            cells0104text = cells0104.add_run("MAT")
            cells0104text.bold = True
            cells[4].width = Cm(2.5)
    
            cells0105 = cells[5].paragraphs[0]
            cells0105text = cells0105.add_run("BIO")
            cells0105text.bold = True
            cells[5].width = Cm(2.5)
    
            cells0106 = cells[6].paragraphs[0]
            cells0106text = cells0106.add_run("PHY")
            cells0106text.bold = True
            cells[6].width = Cm(2.5)
    
            cells0107 = cells[7].paragraphs[0]
            cells0107text = cells0107.add_run("CHE")
            cells0107text.bold = True
            cells[7].width = Cm(2.5)
    
            cells0108 = cells[8].paragraphs[0]
            cells0108text = cells0108.add_run("HIS")
            cells0108text.bold = True
            cells[8].width = Cm(2.5)
    
            cells0109 = cells[9].paragraphs[0]
            cells0109text = cells0109.add_run("GEO")
            cells0109text.bold = True
            cells[9].width = Cm(2.5)
    
            cells0110 = cells[10].paragraphs[0]
            cells0110text = cells0110.add_run("CRE")
            cells0110text.bold = True
            cells[10].width = Cm(2.5)
            
            cells0111 = cells[11].paragraphs[0]
            cells0111text = cells0111.add_run("AGR")
            cells0111text.bold = True
            cells[11].width = Cm(2.5)
    
            cells0112 = cells[12].paragraphs[0]
            cells0112text = cells0112.add_run("COM")
            cells0112text.bold = True
            cells[12].width = Cm(2.5)
            
            cells0113 = cells[13].paragraphs[0]
            cells0113text = cells0113.add_run("BUS")
            cells0113text.bold = True
            cells[13].width = Cm(2.5)
            
            cells0114 = cells[14].paragraphs[0]
            cells0114text = cells0114.add_run("TOT")
            cells0114text.bold = True
            cells[14].width = Cm(2.5)
            
            cells0115 = cells[15].paragraphs[0]
            cells0115text = cells0115.add_run("GRD")
            cells0115text.bold = True
            cells[15].width = Cm(2.5)
            
            cells0116 = cells[16].paragraphs[0]
            cells0116text = cells0116.add_run("PST")
            cells0116text.bold = True
            cells[16].width = Cm(2.5)
         
            pst = 1
            
            for row in student_marks:
                row_cells = table1.add_row().cells
        
                row_cells[0].text = f"{row[0]}"
                row_cells[1].text = f"{(row[1]).upper()}"
                row_cells[2].text = f"{row[2][0]}{row[2][1]}" # Compulsory subject
                row_cells[3].text = f"{row[3][0]}{row[3][1]}" # Compulsory subject
                row_cells[4].text = f"{row[4][0]}{row[4][1]}" # Compulsory subject
                row_cells[5].text = f"{row[5][0]}{row[5][1]}" if row[5][0] != -1 else "-"
                row_cells[6].text = f"{row[6][0]}{row[6][1]}" if row[6][0] != -1 else "-"
                row_cells[7].text = f"{row[7][0]}{row[7][1]}" if row[7][0] != -1 else "-"
                row_cells[8].text = f"{row[8][0]}{row[8][1]}" if row[8][0] != -1 else "-"
                row_cells[9].text = f"{row[9][0]}{row[9][1]}" if row[9][0] != -1 else "-"
                row_cells[10].text = f"{row[10][0]}{row[10][1]}" if row[10][0] != -1 else "-"
                row_cells[11].text = f"{row[11][0]}{row[11][1]}" if row[11][0] != -1 else "-"
                row_cells[12].text = f"{row[12][0]}{row[12][1]}" if row[12][0] != -1 else "-"
                row_cells[13].text = f"{row[13][0]}{row[13][1]}" if row[13][0] != -1 else "-"
                row_cells[14].text = f"{row[14]}"
                row_cells[15].text = f"{row[15]}"
                row_cells[16].text = f"{pst}"
                
                pst += 1
                
                
        # SUBJECT ANALYSIS
        
        doc.add_page_break()
    
        subj_para = doc.add_paragraph()
        subj = subj_para.add_run("SUBJECT ANALYSIS")
        subj.underline = True
        subj.font.color.rgb = RGBColor.from_string("FF0000")
          
          
        # TABLE FOR SUBJECT ANALYSIS
        
        table2 = doc.add_table(rows = 1, cols = 17)
        table2.style = doc.styles["Light Grid Accent 2"]
    
        table2.columns[0].width = Cm(7.0)
        table2.columns[1].width = Cm(3.0)
        table2.columns[2].width = Cm(2.7)
        table2.columns[3].width = Cm(2.7)
        table2.columns[4].width = Cm(2.7)
        table2.columns[5].width = Cm(2.7)
        table2.columns[6].width = Cm(2.7)
        table2.columns[7].width = Cm(2.7)
        table2.columns[8].width = Cm(2.7)
        table2.columns[9].width = Cm(2.7)
        table2.columns[10].width = Cm(2.7)
        table2.columns[11].width = Cm(2.7)
        table2.columns[12].width = Cm(2.7)
        table2.columns[13].width = Cm(2.7)
        table2.columns[14].width = Cm(2.7)
        table2.columns[15].width = Cm(2.7)
        table2.columns[16].width = Cm(2.7)
   
        cells = table2.rows[0].cells
    
        cells0200 = cells[0].paragraphs[0]
        cells0200text = cells0200.add_run("SUBJECTS")
        cells0200text.bold = True
        cells[0].width = Cm(7.0)
    
        cells0201 = cells[1].paragraphs[0]
        cells0201text = cells0201.add_run("ENTRY")
        cells0201text.bold = True
        cells[1].width = Cm(3.0)
        
        cells0202 = cells[2].paragraphs[0]
        cells0202text = cells0202.add_run("A")
        cells0202text.bold = True
        cells[2].width = Cm(2.7)
    
        cells0203 = cells[3].paragraphs[0]
        cells0203text = cells0203.add_run("A-")
        cells0203text.bold = True
        cells[3].width = Cm(2.7)
    
        cells0204 = cells[4].paragraphs[0]
        cells0204text = cells0204.add_run("B+")
        cells0204text.bold = True
        cells[4].width = Cm(2.7)
    
        cells0205 = cells[5].paragraphs[0]
        cells0205text = cells0205.add_run("B")
        cells0205text.bold = True
        cells[5].width = Cm(2.7)
    
        cells0206 = cells[6].paragraphs[0]
        cells0206text = cells0206.add_run("B-")
        cells0206text.bold = True
        cells[6].width = Cm(2.7)
    
        cells0207 = cells[7].paragraphs[0]
        cells0207text = cells0207.add_run("C+")
        cells0207text.bold = True
        cells[7].width = Cm(2.7)
    
        cells0208 = cells[8].paragraphs[0]
        cells0208text = cells0208.add_run("C")
        cells0208text.bold = True
        cells[8].width = Cm(2.7)
    
        cells0209 = cells[9].paragraphs[0]
        cells0209text = cells0209.add_run("C-")
        cells0209text.bold = True
        cells[9].width = Cm(2.7)
    
        cells0210 = cells[10].paragraphs[0]
        cells0210text = cells0210.add_run("D+")
        cells0210text.bold = True
        cells[10].width = Cm(2.7)
    
        cells0211 = cells[11].paragraphs[0]
        cells0211text = cells0211.add_run("D")
        cells0211text.bold = True
        cells[11].width = Cm(2.7)
    
        cells0212 = cells[12].paragraphs[0]
        cells0212text = cells0212.add_run("D-")
        cells0212text.bold = True
        cells[12].width = Cm(2.7)
    
        cells0213 = cells[13].paragraphs[0]
        cells0213text = cells0213.add_run("E")
        cells0213text.bold = True
        cells[13].width = Cm(2.7)
    
        cells0214 = cells[14].paragraphs[0]
        cells0214text = cells0214.add_run("TT")
        cells0214text.bold = True
        cells[14].width = Cm(2.7)
    
        cells0215 = cells[15].paragraphs[0]
        cells0215text = cells0215.add_run("GD")
        cells0215text.bold = True
        cells[15].width = Cm(2.7)
    
        cells0216 = cells[16].paragraphs[0]
        cells0216text = cells0216.add_run("PST")
        cells0216text.bold = True
        cells[16].width = Cm(2.7)
    
        eng = Counter([row[2][1] for row in student_marks])
        eng_entry = len([row[2][0] for row in student_marks if row[2][0] != -1])
        eng_points = round((sum([row[2][2] for row in student_marks]) / eng_entry), 3)
        eng_grade = self.get_mean_grade(eng_points)
        
        kis = Counter([row[3][1] for row in student_marks])
        kis_entry = len([row[3][0] for row in student_marks if row[3][0] != -1])
        kis_points = round((sum([row[3][2] for row in student_marks]) / kis_entry), 3)
        kis_grade = self.get_mean_grade(kis_points)
        
        mat = Counter([row[4][1] for row in student_marks])
        mat_entry = len([row[4][0] for row in student_marks if row[4][0] != -1])
        mat_points = round((sum([row[4][2] for row in student_marks]) / mat_entry), 3)
        mat_grade = self.get_mean_grade(mat_points)
        
        bio = Counter([row[5][1] for row in student_marks])
        bio_entry = len([row[5][0] for row in student_marks if row[5][0] != -1])
        bio_points = round((sum([row[5][2] for row in student_marks]) / bio_entry), 3)
        bio_grade = self.get_mean_grade(bio_points)
        
        phy = Counter([row[6][1] for row in student_marks])
        phy_entry = len([row[6][0] for row in student_marks if row[6][0] != -1])
        phy_points = round((sum([row[6][2] for row in student_marks]) / phy_entry), 3)
        phy_grade = self.get_mean_grade(phy_points)
        
        che = Counter([row[7][1] for row in student_marks])
        che_entry = len([row[7][0] for row in student_marks if row[7][0] != -1])
        che_points = round((sum([row[7][2] for row in student_marks]) / che_entry), 3)
        che_grade = self.get_mean_grade(che_points)
        
        his = Counter([row[8][1] for row in student_marks])
        his_entry = len([row[8][0] for row in student_marks if row[8][0] != -1])
        his_points = round((sum([row[8][2] for row in student_marks]) / his_entry), 3)
        his_grade = self.get_mean_grade(his_points)
        
        geo = Counter([row[9][1] for row in student_marks])
        geo_entry = len([row[9][0] for row in student_marks if row[9][0] != -1])
        geo_points = round((sum([row[9][2] for row in student_marks]) / geo_entry), 3)
        geo_grade = self.get_mean_grade(geo_points)
        
        cre = Counter([row[10][1] for row in student_marks])
        cre_entry = len([row[10][0] for row in student_marks if row[10][0] != -1])
        cre_points = round((sum([row[10][2] for row in student_marks]) / cre_entry), 3)
        cre_grade = self.get_mean_grade(cre_points)
        
        agr = Counter([row[11][1] for row in student_marks])
        agr_entry = len([row[11][0] for row in student_marks if row[11][0] != -1])
        agr_points = round((sum([row[11][2] for row in student_marks]) / agr_entry), 3)
        agr_grade = self.get_mean_grade(agr_points)
        
        com = Counter([row[12][1] for row in student_marks])
        com_entry = len([row[12][0] for row in student_marks if row[12][0] != -1])
        com_points = round((sum([row[12][2] for row in student_marks]) / com_entry), 3)
        com_grade = self.get_mean_grade(com_points)
        
        bus = Counter([row[13][1] for row in student_marks])
        bus_entry = len([row[13][0] for row in student_marks if row[13][0] != -1])
        bus_points = round((sum([row[13][2] for row in student_marks]) / bus_entry), 3)
        bus_grade = self.get_mean_grade(bus_points)
        
        sub_analysis = [
               ["ENGLISH",eng_entry,eng.get("A",0),eng.get("A-",0),eng.get("B+",0),eng.get("B",0),eng.get("B-",0),eng.get("C+",0),eng.get("C",0),eng.get("C-",0),eng.get("D+",0),eng.get("D",0),eng.get("D-",0),eng.get("E",0),eng_points,eng_grade],
               ["KISWAHILI",kis_entry,kis.get("A",0),kis.get("A-",0),kis.get("B+",0),kis.get("B",0),kis.get("B-",0),kis.get("C+",0),kis.get("C",0),kis.get("C-",0),kis.get("D+",0),kis.get("D",0),kis.get("D-",0),kis.get("E",0),kis_points,kis_grade],
               ["MATHEMATICS",mat_entry,mat.get("A",0),mat.get("A-",0),mat.get("B+",0),mat.get("B",0),mat.get("B-",0),mat.get("C+",0),mat.get("C",0),mat.get("C-",0),mat.get("D+",0),mat.get("D",0),mat.get("D-",0),mat.get("E",0),mat_points,mat_grade],
               ["BIOLOGY",bio_entry,bio.get("A",0),bio.get("A-",0),bio.get("B+",0),bio.get("B",0),bio.get("B-",0),bio.get("C+",0),bio.get("C",0),bio.get("C-",0),bio.get("D+",0),bio.get("D",0),bio.get("D-",0),bio.get("E",0),bio_points,bio_grade],
               ["PHYSICS",phy_entry,phy.get("A",0),phy.get("A-",0),phy.get("B+",0),phy.get("B",0),phy.get("B-",0),phy.get("C+",0),phy.get("C",0),phy.get("C-",0),phy.get("D+",0),phy.get("D",0),phy.get("D-",0),phy.get("E",0),phy_points,phy_grade],
               ["CHEMISTRY",che_entry,che.get("A",0),che.get("A-",0),che.get("B+",0),che.get("B",0),che.get("B-",0),che.get("C+",0),che.get("C",0),che.get("C-",0),che.get("D+",0),che.get("D",0),che.get("D-",0),che.get("E",0),che_points,che_grade], 
               ["HISTORY",his_entry,his.get("A",0),his.get("A-",0),his.get("B+",0),his.get("B",0),his.get("B-",0),his.get("C+",0),his.get("C",0),his.get("C-",0),his.get("D+",0),his.get("D",0),his.get("D-",0),his.get("E",0),his_points,his_grade],   
               ["GEOGRAPHY",geo_entry,geo.get("A",0),geo.get("A-",0),geo.get("B+",0),geo.get("B",0),geo.get("B-",0),geo.get("C+",0),geo.get("C",0),geo.get("C-",0),geo.get("D+",0),geo.get("D",0),geo.get("D-",0),geo.get("E",0),geo_points,geo_grade],   
               ["CRE",cre_entry,cre.get("A",0),cre.get("A-",0),cre.get("B+",0),cre.get("B",0),cre.get("B-",0),cre.get("C+",0),cre.get("C",0),cre.get("C-",0),cre.get("D+",0),cre.get("D",0),cre.get("D-",0),cre.get("E",0),cre_points,cre_grade],   
               ["AGRICULTURE",agr_entry,agr.get("A",0),agr.get("A-",0),agr.get("B+",0),agr.get("B",0),agr.get("B-",0),agr.get("C+",0),agr.get("C",0),agr.get("C-",0),agr.get("D+",0),agr.get("D",0),agr.get("D-",0),agr.get("E",0),agr_points,agr_grade],   
               ["COMPUTER",com_entry,com.get("A",0),com.get("A-",0),com.get("B+",0),com.get("B",0),com.get("B-",0),com.get("C+",0),com.get("C",0),com.get("C-",0),com.get("D+",0),com.get("D",0),com.get("D-",0),com.get("E",0),com_points,com_grade],   
               ["BUSINESS",bus_entry,bus.get("A",0),bus.get("A-",0),bus.get("B+",0),bus.get("B",0),bus.get("B-",0),bus.get("C+",0),bus.get("C",0),bus.get("C-",0),bus.get("D+",0),bus.get("D",0),bus.get("D-",0),bus.get("E",0),bus_points,bus_grade],
        ]
    
        sub_analysis = sorted(sub_analysis, reverse=True, key=itemgetter(-2))
    
        pst = 1
    
        for row in sub_analysis:
            row_cells = table2.add_row().cells
        
            row_cells[0].text = f"{row[0]}"
            row_cells[1].text = f"{row[1]}"
            row_cells[2].text = f"{row[2]}"
            row_cells[3].text = f"{row[3]}"
            row_cells[4].text = f"{row[4]}"
            row_cells[5].text = f"{row[5]}"
            row_cells[6].text = f"{row[6]}"
            row_cells[7].text = f"{row[7]}"
            row_cells[8].text = f"{row[8]}"
            row_cells[9].text = f"{row[9]}"
            row_cells[10].text = f"{row[10]}"
            row_cells[11].text = f"{row[11]}"
            row_cells[12].text = f"{row[12]}"
            row_cells[13].text = f"{row[13]}"
            row_cells[14].text = f"{row[14]}"
            row_cells[15].text = f"{row[15]}"
            row_cells[16].text = f"{pst}"
     
            pst += 1
            
        
        doc.add_paragraph()
    
        ovr_para = doc.add_paragraph()
        ovr = ovr_para.add_run("OVERALL ANALYSIS")
        ovr.underline = True
        ovr.font.color.rgb = RGBColor.from_string("FF0000")


        # TABLE 3
    
        table3 = doc.add_table(rows = 1, cols = 14)
        table3.style = doc.styles["Light Grid Accent 3"]
    
        table3.columns[0].width = Cm(7.0)
        table3.columns[1].width = Cm(2.7)
        table3.columns[2].width = Cm(2.7)
        table3.columns[3].width = Cm(2.7)
        table3.columns[4].width = Cm(2.7)
        table3.columns[5].width = Cm(2.7)
        table3.columns[6].width = Cm(2.7)
        table3.columns[7].width = Cm(2.7)
        table3.columns[8].width = Cm(2.7)
        table3.columns[9].width = Cm(2.7)
        table3.columns[10].width = Cm(2.7)
        table3.columns[11].width = Cm(2.7)
        table3.columns[12].width = Cm(2.7)
        table3.columns[13].width = Cm(2.7)
    
        cells = table3.rows[0].cells
    
        cell0300 = cells[0].paragraphs[0]
        cell0300text = cell0300.add_run("ENTRY")
        cell0300text.bold = True
        cells[0].width = Cm(7.0)
        
        cells0301 = cells[1].paragraphs[0]
        cells0301text = cells0301.add_run("A")
        cells0301text.bold = True
        cells[1].width = Cm(2.7)
    
        cells0302 = cells[2].paragraphs[0]
        cells0302text = cells0302.add_run("A-")
        cells0302text.bold = True
        cells[2].width = Cm(2.7)
    
        cells0303 = cells[3].paragraphs[0]
        cells0303text = cells0303.add_run("B+")
        cells0303text.bold = True
        cells[3].width = Cm(2.7)
    
        cells0304 = cells[4].paragraphs[0]
        cells0304text = cells0304.add_run("B")
        cells0304text.bold = True
        cells[4].width = Cm(2.7)
    
        cells0305 = cells[5].paragraphs[0]
        cells0305text = cells0305.add_run("B-")
        cells0305text.bold = True
        cells[5].width = Cm(2.7)
    
        cells0306 = cells[6].paragraphs[0]
        cells0306text = cells0306.add_run("C+")
        cells0306text.bold = True
        cells[6].width = Cm(2.7)
    
        cells0307 = cells[7].paragraphs[0]
        cells0307text = cells0307.add_run("C")
        cells0307text.bold = True
        cells[7].width = Cm(2.7)
    
        cells0308 = cells[8].paragraphs[0]
        cells0308text = cells0308.add_run("C-")
        cells0308text.bold = True
        cells[8].width = Cm(2.7)
    
        cells0309 = cells[9].paragraphs[0]
        cells0309text = cells0309.add_run("D+")
        cells0309text.bold = True
        cells[9].width = Cm(2.7)
    
        cells0310 = cells[10].paragraphs[0]
        cells0310text = cells0310.add_run("D")
        cells0310text.bold = True
        cells[10].width = Cm(2.7)
    
        cells0311 = cells[11].paragraphs[0]
        cells0311text = cells0311.add_run("D-")
        cells0311text.bold = True
        cells[11].width = Cm(2.7)
    
        cells0312 = cells[12].paragraphs[0]
        cells0312text = cells0312.add_run("E")
        cells0312text.bold = True
        cells[12].width = Cm(2.7)
        
        cells0313 = cells[13].paragraphs[0]
        cells0313text = cells0313.add_run("X")
        cells0313text.bold = True
        cells[13].width = Cm(2.7)
    
        overall = Counter([row[15] for row in student_marks])
        
        entry = len(student_marks)
        
        overall_analysis = [
                           [entry,overall.get("A",0),overall.get("A-",0),overall.get("B+",0),overall.get("B",0),overall.get("B-",0),overall.get("C+",0),overall.get("C",0),overall.get("C-",0),overall.get("D+",0),overall.get("D",0),overall.get("D-",0),overall.get("E",0),overall.get("X",0)]          
        ]
    
        for row in overall_analysis:
            row_cells = table3.add_row().cells
        
            row_cells[0].text = f"{row[0]}"
            row_cells[1].text = f"{row[1]}"
            row_cells[2].text = f"{row[2]}"
            row_cells[3].text = f"{row[3]}"
            row_cells[4].text = f"{row[4]}"
            row_cells[5].text = f"{row[5]}"
            row_cells[6].text = f"{row[6]}"
            row_cells[7].text = f"{row[7]}"
            row_cells[8].text = f"{row[8]}"
            row_cells[9].text = f"{row[9]}"
            row_cells[10].text = f"{row[10]}"
            row_cells[11].text = f"{row[11]}"
            row_cells[12].text = f"{row[12]}"
            row_cells[13].text = f"{row[13]}"

        
        # CLASS MEAN SCORE AND GRADE
        
        class_mean_score = round((sum([row[-2] for row in sub_analysis]) / 12), 3)
        class_mean_grade = self.get_mean_grade(class_mean_score)   
        
         
        # WRITE MEAN SCORE TO WORD
        
        doc.add_paragraph()
    
        avg_score_para = doc.add_paragraph()
        avg_score_text = avg_score_para.add_run("MEAN SCORE: ")
        avg_score_text.font.color.rgb = RGBColor.from_string("00008B")
            
        avg_score_mean = avg_score_para.add_run(f" {class_mean_score}")
        avg_score_mean.font.color.rgb = RGBColor.from_string("FF0000")
            
        avg_grade_para = doc.add_paragraph()
        avg_grade_text = avg_grade_para.add_run("MEAN GRADE: ")
        avg_grade_text.font.color.rgb = RGBColor.from_string("00008B")
            
        avg_grade_mean = avg_grade_para.add_run(f" {class_mean_grade}")
        avg_grade_mean.font.color.rgb = RGBColor.from_string("FF0000")
            
        docwriter.save_document(os.path.join(app_path, filename))
        
        show_toast(f"Successfully printed analysis report.")
        
    def get_teachers_comments(self, grade):
        from modules.datamanager import Comment
        
        comments = []
        
        cmt = Comment()
        
        htr = ""
        ctr = ""
        
        for comment in cmt.get_comments():
            comments.append(comment)
            
        if grade == "X":
            comment = [dict["GRADE X"] for dict in comments if "GRADE X" in dict][0]
            htr = comment["head_teacher"]
            ctr = comment["class_teacher"]
            
        if grade == "E":
            comment = [dict["GRADE E"] for dict in comments if "GRADE E" in dict][0]
            htr = comment["head_teacher"]
            ctr = comment["class_teacher"]
            
        elif grade == "D-" or grade == "D" or grade == "D+":
            comment = [dict["GRADE D"] for dict in comments if "GRADE D" in dict][0]
            htr = comment["head_teacher"]
            ctr = comment["class_teacher"]
            
        elif grade == "C-" or grade == "C" or grade == "C+":
            comment = [dict["GRADE C"] for dict in comments if "GRADE C" in dict][0]
            htr = comment["head_teacher"]
            ctr = comment["class_teacher"]
            
        elif grade == "B-" or grade == "B" or grade == "B+":
            comment = [dict["GRADE B"] for dict in comments if "GRADE B" in dict][0]
            htr = comment["head_teacher"]
            ctr = comment["class_teacher"]
            
        elif grade == "A-" or grade == "A" or grade == "A+":
            comment = [dict["GRADE A"] for dict in comments if "GRADE A" in dict][0]
            htr = comment["head_teacher"]
            ctr = comment["class_teacher"]
            
        return htr, ctr
            
    def get_subject_and_code(self, subject):
        sbj = Subjects()
        sbj_data = sbj.get_subjects()
        
        subject = [dict[subject] for dict in sbj_data if subject in dict][0]
      
        code = subject["code"]
        sub_fname = subject["full_name"]
        
        return code, sub_fname
        
    def get_student_remark(self, grade):
        grading = Grading()
        grd_data = grading.get_grading_data()
        
        grading_data = [dict["GENERAL"] for dict in grd_data if "GENERAL" in dict][0]
        
        for k,v in grading_data.items():
            if k == grade:
                return v[2]
                
            else:
                continue
                
        return " "
        
    def get_teacher_sign(self, form, grade, subj):
        subject_data = []
        
        sobj = Subjects()
        
        for subject in sobj.get_subjects():
            subject_data.append(subject)
            
        if form == "FORM 1":
            form = "F1"
            
        elif form == "FORM 2":
            form = "F2"
            
        elif form == "FORM 3":
            form = "F3"
            
        elif form == "FORM 4":
            form = "F4"
         
        teacher = [data[subj]["teacher"][form] for data in subject_data if subj in data][0]
        teacher = teacher.split(" ")[0]
        
        if not grade == "" or grade == "X" or grade == " ":
            return teacher
            
        return " "
        
    def get_grade_and_points(self, grading_data, value):
        for k, v in grading_data.items():
            marks_range = [x for x in range(v[0][0], (v[0][1] + 1))]
            
            grade = k
            points = v[1]
        
            if value in marks_range:
                return grade, points
                
            else:
                continue
                
        return "", 0
        
    def get_mean_grade(self, value):
        '''Takes floating point value and returns Grade. Eg. 2.561 will return D-'''
        
        value = int(value)
        
        grading = Grading()
        grd_data = grading.get_grading_data()
        
        grading_data = [dict["GENERAL"] for dict in grd_data if "GENERAL" in dict][0]
        
        for k,v in grading_data.items():
            if v[1] == value:
                return k
                
            else:
                continue
                
        return "X"
       
    def wash_value(self, value):
        if re.match("^[\d]+$", value):
            clean_value = int(value)
            
        else:
            clean_value = -1
        
        return clean_value
        
    def get_one_best_point_among_two(self, eng, kis):
        if eng >= kis:
            return eng
            
        else:
            return kis
            
    def get_five_best_points_among_many(self, *args):
        subjects = sorted(*args, reverse=True)
        best5 = subjects[0:5]
        
        return best5
        
    def get_one_best_subject_eng_or_kisw(self, eng, kis):
        if eng >= kis:
            return eng
            
        else:
            return kis
            
    def get_five_best_subjects(self, *args):
        subjects = sorted(*args, reverse=True)
        best5 = subjects[0:5]
        
        return best5
       
    def get_f1f2_mean_grade(self, value):
        grading = Grading()
        grd_data = grading.get_grading_data()
        
        grading_data = [dict["SCIENCE"] for dict in grd_data if "SCIENCE" in dict][0]
        
        for k, v in grading_data.items():
            marks_range = [x for x in range(v[0][0], (v[0][1] + 1))]
            
            grade = k
            
            if value in marks_range:
                return grade
                
            else:
                continue
                
        return ""
        
    def get_f3f4_mean_grade(self, value):
        grading = Grading()
        grd_data = grading.get_grading_data()
        
        grading_data = [dict["GENERAL"] for dict in grd_data if "GENERAL" in dict][0]
        
        for k, v in grading_data.items():
            marks_range = [x for x in range(v[0][0], (v[0][1] + 1))]
            
            grade = k
            
            if value in marks_range:
                return grade
                
            else:
                continue
                
        return ""
        
    