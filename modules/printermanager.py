import docx, re, os, time
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT

from modules.datamanager import Setting

from datetime import datetime


class SchoolDocument:
    def __init__(self, fontname, fontsize):
        ### INIT DOCUMENT STYLE ###
        self.doc = docx.Document()
        
        # SET SIZE TO A4
    
        self.section = self.doc.sections[0]
        self.section.page_height = Mm(297)
        self.section.page_width = Mm(210)
        self.section.header_distance = Mm(5)
        self.section.footer_distance = Mm(5)
        
        style = self.doc.styles["Normal"]
        font = style.font
        font.size = Pt(fontsize)
        font.name = fontname
        
    def set_portrait_mode(self):
        self.section.top_margin = Inches(0.3)
        self.section.bottom_margin = Inches(0.1)
        
    def set_landscape_mode(self):
        # SET SIZE TO LANDSCAPE MODE
    
        self.section = self.doc.sections[0]
        self.section.orientation = WD_ORIENT.LANDSCAPE
        self.section.page_width = Mm(330)
        self.section.page_height = Mm(220)
        self.section.header_distance = Mm(5)
        self.section.footer_distance = Mm(5)
        
        self.section.top_margin = Inches(0.3)
        self.section.bottom_margin = Inches(0.1)
        
    def write_headers_and_footers(self):
        # ADD CURRENT DATE TO HEADER
        
        logo = "./assets/imgs/logo.png"
        time_now = str(datetime.now().strftime("%A, %d %B %Y"))
  
        header = self.section.header
        header_para = header.add_paragraph()
        header_para_run = header_para.add_run(time_now)
        header_para_run.font.size = Pt(8)
        header_para_run.font.name = "Times New Roman"
        header_para_run.italic = True
        
        # ADD SOFTWARE LOGO TO FOOTER
        
        footer = self.section.footer
        footer_para = footer.add_paragraph()
        footer_para_run = footer_para.add_run().add_picture(logo, width=Cm(2), height=Cm(0.8))
        footer_para.paragraph_format.space_after = Pt(0)
        footer_para.paragraph_format.space_before = Pt(0)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
     
    def write_headings(self):
        # SETTINGS MANAGER #
        
        setting_obj = Setting()
        settings = setting_obj.get_settings()
        
        try:
            schnmecolor = str((settings["school_name_color"]).split("#")[1])
            schcapcolor = str((settings["school_caption_color"]).split("#")[1])
        
        except:
            schnmecolor = "000000"
            schcapcolor = "000000"
        
        school_name = settings["school_name"]
        school_caption = settings["school_caption"]
        school_address = settings["school_address"]
        
        # # # WRITE HEADINGs # # #
        
        heading = self.doc.add_paragraph()
        
        heading1 = heading.add_run(f"{school_name.upper()}\n")
        heading1.font.size = Pt(22)
        heading1.bold = True
        heading1.font.name = "Times New Roman"
        heading1.font.color.rgb = RGBColor.from_string(schnmecolor)
          
        heading2 = heading.add_run(f"{school_caption.upper()}\n")
        heading2.font.size = Pt(15)
        heading2.font.name = "Times New Roman"
        heading2.font.color.rgb = RGBColor.from_string(schcapcolor)
        heading2.bold = True
            
        heading3 = heading.add_run(f"{school_address.upper()}\n")
        heading3.font.size = Pt(12)
        heading3.font.name = "Times New Roman"
        heading3.font.color.rgb = RGBColor.from_string("000000")
        heading3.bold = True
        
        heading.paragraph_format.line_spacing = Pt(20)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.alignment = 1
        
    def write_title(self, dtitle, space_after=15):    
        title = self.doc.add_paragraph()
        
        title_run = title.add_run(dtitle)
        title_run.font.name = "Impact"
        title_run.font.size = Pt(15)
        title_run.font.color.rgb = RGBColor.from_string("0047AB")
        title_run.underline = True
        
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(space_after)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def write_student_report_card_heading_details(self, student_name, adm, form, term, exam, year, grade, position, out_of):
        ##### STUDENT DETAILS PARAGRAPH #####
        
        dot_color = "FFFFFF"
        
        name_para = self.doc.add_paragraph()
        paragraph_format = name_para.paragraph_format
        paragraph_format.line_spacing = Pt(13)

        name_para.add_run("NAME ")

        std_name = name_para.add_run(f"     {student_name.upper()}      ")
        std_name.underline = True
        std_name.underline = WD_UNDERLINE.DOTTED
        std_name.bold = True
        std_name.font.name = "Lucida Calligraphy"

        name_para.add_run("ADM NO. ")

        adm_no = name_para.add_run(f"      {adm}      ")
        adm_no.underline = True
        adm_no.underline = WD_UNDERLINE.DOTTED
        adm_no.bold = True
        adm_no.font.name = "Lucida Calligraphy"

        name_para.add_run("FORM ")

        form = name_para.add_run(f"      {form}     ")
        form.underline = True
        form.underline = WD_UNDERLINE.DOTTED
        form.bold = True

        name_para.add_run("TERM ")

        term = name_para.add_run(f"     {term}     ")
        term.underline = True
        term.underline = WD_UNDERLINE.DOTTED
        term.bold = True
        term.font.name = "Lucida Calligraphy"

        name_para.add_run("EXAM")

        exam = name_para.add_run(f"     {exam}     ")
        exam.underline = True
        exam.underline = WD_UNDERLINE.DOTTED
        exam.bold = True
        exam.font.name = "Lucida Calligraphy"

        # YEAR
        
        name_para.add_run("YEAR ")

        year = name_para.add_run(f"     {year}     ")
        
        ct_dot_run = name_para.add_run("_")
        ct_dot_run.font.color.rgb = RGBColor.from_string(dot_color)
            
        year.underline = True
        year.underline = WD_UNDERLINE.DOTTED
        year.bold = True
        year.font.name = "Lucida Calligraphy"
       
        ##### GRADE AND POSITION PARAGRAPH #####
        
        position_para = self.doc.add_paragraph()
        paragraph_format = position_para.paragraph_format
        paragraph_format.line_spacing = Pt(15)
        
        # GRADE #
        
        position_para.add_run("MEAN GRADE ")

        mgrade = position_para.add_run(f"         {grade}        ")
        mgrade.underline = True
        mgrade.underline = WD_UNDERLINE.DOTTED
        mgrade.bold = True
        mgrade.font.name = "Lucida Calligraphy"
        
        # POSITION #
        
        position_para.add_run("POSITION")
        
        pstn = position_para.add_run(f"          {position}          ")
        pstn.underline = True
        pstn.underline = WD_UNDERLINE.DOTTED
        pstn.bold = True
        pstn.font.name = "Lucida Calligraphy"
        
        # OUT OF #
        
        position_para.add_run("OUT OF")
        
        outof = position_para.add_run(f"          {out_of}          ")
        
        ct_dot_run = position_para.add_run("_")
        ct_dot_run.font.color.rgb = RGBColor.from_string(dot_color)
            
        outof.underline = True
        outof.underline = WD_UNDERLINE.DOTTED
        outof.bold = True
        outof.font.name = "Lucida Calligraphy"

    def write_next(self):
        return self.doc
        
    def save_document(self, fname):
        self.doc.save(fname)
